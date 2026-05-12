import io
from docx import Document
import csv

import pandas as pd
import pdfplumber
import json
import os
import concurrent.futures

RUIDO = {
    "nan",
    "none",
    "0.0",
    "0",
    "",
    "n/a",
    "tbd",
    "tbc",
    "-",
    "--",
    "---",
    "#n/a",
    "#ref!",
    "#value!",
    "#name?",
}


def _detectar_separador_csv(conteudo: str) -> str:
    """Auto-detecta o separador do CSV (,, ;, \\t, |)."""
    # Tomar primeiras linhas para análise
    linhas = conteudo.split('\n')[:5]
    
    separadores = {',': 0, ';': 0, '\t': 0, '|': 0}
    
    for linha in linhas:
        if not linha.strip():
            continue
        for sep in separadores:
            separadores[sep] += linha.count(sep)
    
    # Retornar o separador mais comum (deve ter contagem > 0)
    mais_comum = max(separadores.items(), key=lambda x: x[1])
    return mais_comum[0] if mais_comum[1] > 0 else ','


def read_document(file) -> tuple[str, list]:
    """Lê Excel, CSV ou PDF. Devolve (texto, paginas_sem_texto)."""
    paginas_sem_texto = []

    if file.name.lower().endswith(".pdf"):
        partes = []
        with pdfplumber.open(io.BytesIO(file.read())) as pdf:
            for i, p in enumerate(pdf.pages):
                texto_pag = p.extract_text(layout=True)
                if texto_pag:
                    partes.append(f"[Pág: {i+1}] {texto_pag}")
                else:
                    paginas_sem_texto.append(i + 1)
        return "\n".join(partes), paginas_sem_texto
    if file.name.lower().endswith(('.docx')):
        # Lógica para ler documentos .docx  
        partes = []
        try:
            doc = Document(io.BytesIO(file.read()))
            
            # Extrair parágrafos
            para_num = 0
            for para in doc.paragraphs:
                texto = para.text.strip()
                if texto:  # Ignorar parágrafos vazios
                    para_num += 1
                    partes.append(f"[Parágrafo: {para_num}] {texto}")
            
            # Extrair tabelas
            if doc.tables:
                partes.append("\n[TABELAS DO DOCUMENTO]\n")
                for table_idx, table in enumerate(doc.tables, 1):
                    partes.append(f"\n[Tabela: {table_idx}]")
                    for row_idx, row in enumerate(table.rows, 1):
                        cells_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text.lower() not in RUIDO and cell_text:
                                cells_text.append(cell_text)
                        if cells_text:
                            partes.append(f"  [Linha {row_idx}] {' | '.join(cells_text)}")
            
            return "\n".join(partes), []
        except Exception as e:
            return f"[Erro a ler DOCX: {str(e)}]", []
    elif file.name.lower().endswith('.csv'):
        # Lógica para ler ficheiros CSV com auto-detect de separador
        lines = []
        conteudo_bytes = file.read()
        
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
            try:
                conteudo_texto = conteudo_bytes.decode(encoding)
                separador = _detectar_separador_csv(conteudo_texto)
                
                df = pd.read_csv(
                    io.BytesIO(conteudo_bytes), 
                    encoding=encoding,
                    sep=separador,
                    on_bad_lines='skip',  # Ignorar linhas mal formatadas
                    dtype=str,  # Ler tudo como string inicialmente
                )
                
                for idx, row in df.iterrows():
                    vals = []
                    for v in row:
                        if pd.isna(v) or v is None:
                            continue
                        cell_text = str(v).strip()
                        if not cell_text:
                            continue
                        if cell_text.lower() not in RUIDO and len(cell_text) > 1:
                            vals.append(cell_text)
                    if vals:  # Se tem pelo menos 1 valor
                        lines.append(f"[Linha: {idx+2}] {' | '.join(vals)}")
                
                # Se conseguiu ler, retorna
                return "\n".join(lines) if lines else "[CSV vazio ou sem dados válidos]", []
                
            except UnicodeDecodeError:
                continue
            except pd.errors.ParserError as e:
                # Tentar com mais separadores ou skip de linhas problemáticas
                continue
            except Exception as e:
                continue
        
        # Se chegou aqui, todos os encodings falharam
        return f"[Erro: Não foi possível ler o ficheiro CSV com encoding UTF-8, Latin-1, CP1252 ou ISO-8859-1]", []
    else:
        # Lógica para ler ficheiros Excel
        xls = pd.ExcelFile(file)
        lines = []
        for sheet in xls.sheet_names:
            df = xls.parse(sheet)
            for idx, row in df.iterrows():
                vals = []
                for v in row:
                    if pd.isna(v):
                        continue
                    cell_text = str(v).strip()
                    if not cell_text:
                        continue
                    if cell_text.lower() not in RUIDO and len(cell_text) > 1:
                        vals.append(cell_text)
                if len(vals) > 1:
                    lines.append(f"[Linha: {idx+2}] {' | '.join(vals)}")
        return "\n".join(lines), []

# Carregar o JSON
def carregar_regras_json():
    caminho_json = os.path.join(os.path.dirname(os.path.dirname(__file__)), "RegrasMekkin.json")
    
    fallback = {
        "regra_final": "Assumir que toda a informacao e IRRELEVANTE ate demonstrar impacto na estrutura.",
        "ignorar_estritamente": ["Arquitetura", "Cores", "AVAC sem carga", "Betão sem interface"]
    }
    
    if os.path.exists(caminho_json):
        with open(caminho_json, 'r', encoding='utf-8') as f:
            return json.dumps(json.load(f), ensure_ascii=False, indent=2)
    return json.dumps(fallback, ensure_ascii=False, indent=2)

